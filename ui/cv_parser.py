"""
Lectura y analisis de CV para la UI, 100% local, por reglas/keywords.

Sin APIs externas, sin IA externa. El archivo se lee EN MEMORIA (bytes de
st.file_uploader via BytesIO) y nunca se escribe a disco desde aca -- si
algun consumidor necesita persistirlo temporalmente, es su responsabilidad
usar una carpeta ignorada por git (ver .gitignore: uploads/, data/cvs/) y
borrarla despues.

Reusa core.constants.CATEGORIAS_KEYWORDS para que las categorias sugeridas
sean las mismas que ya usa el resto de la app (categoria_detectada, filtro
de categorias en la sidebar) -- no se inventa un vocabulario nuevo.
"""
import io
import re

from core import constants

# ----------------------------------------------------------------------
# LECTURA DE ARCHIVO (en memoria, sin tocar disco)
# ----------------------------------------------------------------------

def leer_texto_cv(uploaded_file):
    """Recibe el objeto de st.file_uploader y devuelve el texto extraido.
    Despacha por extension del nombre de archivo. Lanza ValueError si el
    formato no es PDF/DOCX/TXT (la UI debe mostrar ese error, no crashear)."""
    nombre = (uploaded_file.name or "").lower()
    datos = uploaded_file.getvalue()  # bytes en memoria, no se guarda nada

    if nombre.endswith(".pdf"):
        texto = _leer_pdf(datos)
    elif nombre.endswith(".docx"):
        texto = _leer_docx(datos)
    elif nombre.endswith(".txt"):
        texto = _leer_txt(datos)
    else:
        raise ValueError(f"Formato no soportado: '{uploaded_file.name}'. Usar PDF, DOCX o TXT.")
    return _limpiar_texto(texto)


# Caracter de reemplazo (U+FFFD, lo que queda cuando no se puede decodificar
# un byte/glyph) + zona de "uso privado" Unicode (ahi caen la mayoria de los
# bullets/iconos de fuentes custom que los PDFs embeben y pypdf no puede
# traducir a texto real) + caracteres de control sueltos que a veces deja
# una extraccion de PDF rota. Se borran sin intentar "adivinar" el simbolo
# original -- no hay forma confiable de recuperarlo sin OCR/IA.
_PATRON_CARACTER_INVALIDO = re.compile(r"[�-\x00-\x08\x0b\x0c\x0e-\x1f]")


def _limpiar_texto(texto):
    """Saca caracteres invalidos/bullets rotos (tipicos de PDFs con fuentes
    de icono custom) y el espacio en blanco de mas que dejan al borrarse."""
    if not texto:
        return texto
    texto = _PATRON_CARACTER_INVALIDO.sub("", texto)
    lineas = []
    for linea in texto.splitlines():
        limpia = re.sub(r"^[ \t]+", "", linea)
        limpia = re.sub(r"[ \t]{2,}", " ", limpia).rstrip()
        lineas.append(limpia)
    return "\n".join(lineas)


def _leer_pdf(datos):
    from pypdf import PdfReader
    lector = PdfReader(io.BytesIO(datos))
    paginas = [(pagina.extract_text() or "") for pagina in lector.pages]
    texto = "\n".join(paginas)
    if not texto.strip():
        # PDF escaneado/imagen sin capa de texto: pypdf no hace OCR. Se
        # avisa como texto vacio, la UI decide como mostrar el aviso.
        return ""
    return texto


def _leer_docx(datos):
    import docx
    documento = docx.Document(io.BytesIO(datos))
    partes = [p.text for p in documento.paragraphs]
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    partes.append(celda.text)
    return "\n".join(partes)


def _leer_txt(datos):
    for codificacion in ("utf-8", "latin-1"):
        try:
            return datos.decode(codificacion)
        except UnicodeDecodeError:
            continue
    return datos.decode("utf-8", errors="replace")


# ----------------------------------------------------------------------
# VOCABULARIOS (reglas simples, sin ML)
# ----------------------------------------------------------------------

SKILLS_TECNICOS = [
    "python", "sql", "power bi", "power query", "vba", "excel avanzado",
    "excel", "html", "css", "javascript", "typescript", "java", "c#",
    "c++", ".net", "php", "react", "node", "angular", "django", "flask",
    "aws", "azure", "gcp", "docker", "kubernetes", "linux", "bash",
    "machine learning", "sap", "erp", "selenium", "git", "github",
    "tableau", "looker", "oracle", "mysql", "postgresql", "mongodb",
    "salesforce", "r studio", "spss",
]

HERRAMIENTAS = [
    "excel", "power bi", "sap", "jira", "git", "github", "office",
    "word", "outlook", "trello", "slack", "salesforce", "tableau",
    "looker", "power query", "sql server", "servicenow", "zendesk",
    "notion", "confluence", "asana",
]

IDIOMAS_CONOCIDOS = [
    "ingles", "inglés", "español", "espanol", "italiano", "portugues",
    "portugués", "frances", "francés", "aleman", "alemán", "chino",
    "mandarin", "mandarín",
]

_UBICACIONES_CONOCIDAS = [
    ("ciudad autonoma de buenos aires", "CABA"),
    ("ciudad autónoma de buenos aires", "CABA"),
    ("capital federal", "CABA"),
    ("caba", "CABA"),
    ("gran buenos aires", "GBA"),
    ("gba", "GBA"),
    ("buenos aires", "Buenos Aires"),
    ("la plata", "La Plata"),
    ("cordoba", "Córdoba"),
    ("córdoba", "Córdoba"),
    ("rosario", "Rosario"),
    ("mendoza", "Mendoza"),
    ("remoto", "Remoto"),
]

# Categoria -> 1 termino "nucleo" para sugerir como palabra obligatoria.
# A PROPOSITO un solo termino por categoria (no toda la bolsa de skills):
# el filtro "palabras obligatorias" de la UI exige que TODAS esten en el
# titulo (semantica AND) -- una sugerencia con 5-8 skills dejaria 0
# resultados casi siempre. Ver buscador_core._aplicar_filtros_demo.
CATEGORIA_KEYWORD_CORE = {
    "Data / BI": "datos",
    "Soporte IT": "soporte",
    "QA / Testing": "qa",
    "Desarrollo": "desarrollador",
    "Analista Funcional": "funcional",
    "Supply Chain": "supply",
    "Administrativo / Procesos": "administrativo",
    "Técnico / Producción": "tecnico",
}

# Categoria -> puesto objetivo sugerido (texto amigable, solo informativo).
CATEGORIA_A_PUESTO = {
    "Data / BI": "Analista de Datos Junior",
    "Soporte IT": "Soporte IT Junior",
    "QA / Testing": "QA Tester Junior",
    "Desarrollo": "Desarrollador Junior",
    "Analista Funcional": "Analista Funcional Junior",
    "Supply Chain": "Analista de Supply Chain",
    "Administrativo / Procesos": "Administrativo con perfil técnico",
    "Técnico / Producción": "Técnico / Producción (fuera de foco IT/Data)",
}

# core.constants.CATEGORIAS_KEYWORDS puntua por frases especificas (ej. "power bi",
# "python developer") y no por skills sueltas -- un CV que solo dice
# "Python, SQL, Power BI, ERP" en una lista de habilidades no suma nada a
# Data/BI o Analista Funcional con esas listas solas. Este bono conecta las
# skills que ya extrae _extraer_skills() con la categoria IT/Data que
# realmente corresponden, para que no queden opacadas por texto narrativo
# de otro rubro (ver _SKILLS_IT_FUERTES mas abajo).
_BONUS_CATEGORIA_POR_SKILL = {
    "python": {"Data / BI": 1, "Desarrollo": 1},
    "sql": {"Data / BI": 2},
    "power bi": {"Data / BI": 2},
    "power query": {"Data / BI": 1},
    "oracle": {"Data / BI": 1},
    "erp": {"Analista Funcional": 2},
    "sap": {"Analista Funcional": 2},
    "git": {"Desarrollo": 1},
    "html": {"Desarrollo": 1},
    "css": {"Desarrollo": 1},
    "javascript": {"Desarrollo": 2},
    "java": {"Desarrollo": 1},
    "selenium": {"QA / Testing": 2},
    "jira": {"QA / Testing": 1, "Analista Funcional": 1},
}

# Si aparece cualquiera de estas skills, el CV tiene señal tecnica/IT real
# -- en ese caso Supply Chain no puede ganarle a las categorias IT solo
# porque el relato de experiencia previa (compras/abastecimiento/inventario)
# tiene mas frases sueltas que matchean. Pedido explicito: "que Supply
# Chain no opaque categorias IT cuando hay skills como SQL, Python, Power
# BI, ERP, Oracle, Git, HTML/CSS/JS".
_SKILLS_IT_FUERTES = {
    "python", "sql", "power bi", "power query", "erp", "sap", "oracle",
    "git", "html", "css", "javascript", "java", "c#", "c++", ".net",
    "react", "node", "azure", "aws", "gcp", "docker", "selenium",
}

_ENCABEZADOS_SECCION = {
    "experiencia": [
        "experiencia", "experiencia laboral", "experiencia profesional",
        "trayectoria laboral", "historial laboral", "antecedentes laborales",
    ],
    "educacion": [
        "educacion", "educación", "formacion academica", "formación académica",
        "formacion", "formación", "estudios",
    ],
}
_TODOS_LOS_ENCABEZADOS = {
    e for lista in _ENCABEZADOS_SECCION.values() for e in lista
} | {
    "idiomas", "languages", "habilidades", "skills", "conocimientos",
    "competencias", "referencias", "contacto", "datos personales",
}

_PATRON_EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PATRON_TELEFONO = re.compile(
    r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?)?\d{3,4}[\s.-]?\d{3,4}"
)
_PATRON_NOMBRE = re.compile(
    r"^[A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ'-]+(\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ'-]+){1,3}$"
)
_PALABRAS_NO_NOMBRE = {"curriculum vitae", "curriculum", "cv", "resume", "perfil profesional"}


# ----------------------------------------------------------------------
# EXTRACCION (funciones chicas, cada una hace 1 cosa)
# ----------------------------------------------------------------------

def _extraer_email(texto):
    m = _PATRON_EMAIL.search(texto)
    return m.group(0) if m else ""


def _extraer_telefono(texto):
    for m in _PATRON_TELEFONO.finditer(texto):
        digitos = re.sub(r"\D", "", m.group(0))
        if 8 <= len(digitos) <= 13:
            return m.group(0).strip()
    return ""


def _extraer_nombre(texto):
    """Heuristica simple: primera linea de las primeras 8 que parezca un
    nombre propio (2-4 palabras capitalizadas, sin digitos ni @). No hay
    forma 100% confiable sin NLP/IA -- se deja siempre editable en la UI."""
    for linea in texto.splitlines()[:8]:
        candidato = linea.strip()
        if not candidato or len(candidato) > 60:
            continue
        if candidato.lower().rstrip(":") in _PALABRAS_NO_NOMBRE:
            continue
        if "@" in candidato or any(ch.isdigit() for ch in candidato):
            continue
        if _PATRON_NOMBRE.match(candidato):
            return candidato
    return ""


def _extraer_ubicacion(texto_lower):
    for patron, etiqueta in _UBICACIONES_CONOCIDAS:
        if patron in texto_lower:
            return etiqueta
    return ""


def _detectar_seniority(texto_lower):
    """Orden de chequeo mas especifico primero (semi senior antes que
    senior suelto). Riesgo residual conocido: 'sr.'/'ssr' tambien podrian
    ser abreviatura de 'señor', se prioriza no perderse la seniority real
    de un CV."""
    if re.search(r"\bsemi[-\s]?senior\b|\bsemi\s+s?sr\.?\b", texto_lower):
        return "Semi Senior"
    if re.search(r"\bsenior\b|\bs?sr\.?\b", texto_lower):
        return "Senior"
    if re.search(r"\btrainee\b", texto_lower):
        return "Trainee"
    if re.search(r"\bjunior\b|\bjr\.?\b", texto_lower):
        return "Junior"
    return "No detectado"


def _extraer_skills(texto_lower):
    return [kw for kw in SKILLS_TECNICOS if kw in texto_lower]


def _extraer_herramientas(texto_lower, skills_ya_detectadas):
    ya = set(skills_ya_detectadas)
    return [kw for kw in HERRAMIENTAS if kw in texto_lower and kw not in ya]


def _extraer_idiomas(texto_lower):
    encontrados = []
    for idioma in IDIOMAS_CONOCIDOS:
        if idioma in texto_lower:
            etiqueta = idioma.capitalize()
            if etiqueta not in encontrados and idioma.replace("é", "e").replace("á", "a") not in [
                i.lower().replace("é", "e").replace("á", "a") for i in encontrados
            ]:
                encontrados.append(etiqueta)
    return encontrados


def _extraer_seccion(lineas, encabezados):
    """Busca una linea que sea (o empiece con) alguno de 'encabezados' y
    devuelve el bloque de texto hasta el proximo encabezado conocido de
    cualquier seccion, o hasta el final. Recorte de 1200 caracteres para
    no inundar la pantalla con un CV entero."""
    idx_inicio = None
    for i, linea in enumerate(lineas):
        limpio = linea.strip().lower().rstrip(":")
        if limpio in encabezados:
            idx_inicio = i + 1
            break
    if idx_inicio is None:
        return ""

    idx_fin = len(lineas)
    for j in range(idx_inicio, len(lineas)):
        limpio = lineas[j].strip().lower().rstrip(":")
        if limpio and limpio in _TODOS_LOS_ENCABEZADOS:
            idx_fin = j
            break

    bloque = "\n".join(l for l in lineas[idx_inicio:idx_fin] if l.strip())
    return bloque.strip()[:1200]


def _sugerir_categorias(texto_lower, skills, top_n=3):
    """Reusa core.constants.CATEGORIAS_KEYWORDS (las mismas listas que usa
    core.classifier.detectar_categoria() para las ofertas) para puntuar el
    CV completo, mas el bono de _BONUS_CATEGORIA_POR_SKILL para que las
    skills tecnicas concretas (no solo frases narrativas) cuenten.

    Supply Chain se trata aparte a proposito: si el CV tiene skills IT
    fuertes (_SKILLS_IT_FUERTES), su score se penaliza y ademas SIEMPRE
    se devuelve al final de la lista (secundaria), nunca como principal
    -- aunque el score crudo le hubiera ganado a alguna categoria IT por
    tener mas frases sueltas de "compras/abastecimiento/inventario" en el
    relato de experiencia previa."""
    scores = {}
    for categoria, keywords in constants.CATEGORIAS_KEYWORDS.items():
        cantidad = sum(1 for kw in keywords if kw in texto_lower)
        if cantidad > 0:
            scores[categoria] = cantidad

    for skill in skills:
        for categoria, bono in _BONUS_CATEGORIA_POR_SKILL.get(skill, {}).items():
            scores[categoria] = scores.get(categoria, 0) + bono

    supply_score = scores.pop("Supply Chain", 0)
    if any(s in skills for s in _SKILLS_IT_FUERTES):
        supply_score *= 0.4

    ordenadas = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
    sugeridas = [categoria for categoria, _ in ordenadas[:top_n]]

    if supply_score > 0 and "Supply Chain" not in sugeridas:
        sugeridas.append("Supply Chain")

    return sugeridas


def _sugerir_seniority_objetivo(seniority_detectada):
    """Nivel para BUSCAR, no necesariamente el nivel real del CV -- esta
    herramienta busca a proposito roles Junior/Trainee/Semi Senior
    (core.classifier.clasificar_decision descarta Senior/SSr directo),
    asi que "Senior" o "No detectado" en el CV (tipico de alguien
    migrando de otro rubro a IT/Data, o de un CV sin señal clara) no
    tienen sentido como objetivo de busqueda -- se sugiere el nivel de
    entrada. Semi Senior/Trainee detectados se respetan tal cual. Siempre
    editable en pantalla."""
    if seniority_detectada == "Semi Senior":
        return ["Semi Senior"]
    if seniority_detectada == "Trainee":
        return ["Trainee"]
    return ["Junior", "Trainee"]


# ----------------------------------------------------------------------
# ENTRADA PRINCIPAL
# ----------------------------------------------------------------------

def analizar_cv(texto):
    """Punto de entrada unico: recibe el texto ya extraido (de
    leer_texto_cv) y devuelve un dict con el perfil detectado. Todo por
    reglas/keywords locales -- nada de aca llama a una API externa."""
    texto = texto or ""
    texto_lower = texto.lower()
    lineas = texto.splitlines()

    skills = _extraer_skills(texto_lower)
    categorias_sugeridas = _sugerir_categorias(texto_lower, skills)
    seniority_detectada = _detectar_seniority(texto_lower)
    seniority_objetivo = _sugerir_seniority_objetivo(seniority_detectada)

    puestos_objetivo = []
    for c in categorias_sugeridas:
        etiqueta = CATEGORIA_A_PUESTO.get(c)
        if not etiqueta:
            continue
        # Si el CV menciona SAP/ERP, el puesto de Analista Funcional se
        # afina -- es el mismo perfil que "Soporte Funcional ERP/SAP" en
        # las busquedas reales (ver BUSQUEDAS_ERP_SAP en el script).
        if c == "Analista Funcional" and any(s in skills for s in ("sap", "erp")):
            etiqueta = "Analista Funcional / Soporte ERP-SAP Junior"
        puestos_objetivo.append(etiqueta)

    # Suave a proposito: NO se auto-aplica ninguna palabra obligatoria por
    # default (una sola palabra obligatoria, ej. "datos", puede dejar
    # afuera ofertas buenas de Soporte IT/Analista Funcional en perfiles
    # mixtos). Esto queda como SUGERENCIA opcional -- la UI la muestra
    # aparte y solo la carga en el filtro si el usuario confirma con un
    # checkbox. Se sigue limitando a 1 termino (de la categoria principal,
    # nunca Supply Chain si hay algo mejor) por la semantica AND del
    # filtro real -- ver CATEGORIA_KEYWORD_CORE.
    palabras_incluir = []
    categoria_principal = next((c for c in categorias_sugeridas if c != "Supply Chain"), None)
    categoria_principal = categoria_principal or (categorias_sugeridas[0] if categorias_sugeridas else None)
    if categoria_principal:
        core_term = CATEGORIA_KEYWORD_CORE.get(categoria_principal)
        if core_term:
            palabras_incluir = [core_term]

    # Estas SI se auto-aplican (solo excluyen, nunca pueden vaciar
    # resultados por AND) -- salvo que el objetivo incluya Semi Senior,
    # en cuyo caso no tiene sentido excluir "semi senior"/"ssr".
    palabras_excluir = []
    if "Semi Senior" not in seniority_objetivo:
        palabras_excluir = ["senior", "semi senior", "semi-senior", "ssr", "sr"]

    return {
        "nombre": _extraer_nombre(texto),
        "email": _extraer_email(texto),
        "telefono": _extraer_telefono(texto),
        "ubicacion": _extraer_ubicacion(texto_lower),
        "skills": skills,
        "herramientas": _extraer_herramientas(texto_lower, skills),
        "idiomas": _extraer_idiomas(texto_lower),
        "experiencia": _extraer_seccion(lineas, _ENCABEZADOS_SECCION["experiencia"]),
        "educacion": _extraer_seccion(lineas, _ENCABEZADOS_SECCION["educacion"]),
        "seniority_detectada": seniority_detectada,
        "seniority_objetivo": seniority_objetivo,
        "categorias_sugeridas": categorias_sugeridas,
        "puestos_objetivo": puestos_objetivo,
        "palabras_clave_incluir": palabras_incluir,
        "palabras_clave_excluir": palabras_excluir,
    }
